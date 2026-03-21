"""
stress_test_gen.py
Generate 4 test DOCX files for the spelling-only reviewer stress test.
Run: python tests/stress_test_gen.py

Outputs (in appraisal-automation/_temp/):
  test_spelling_typeA.docx
  test_spelling_typeB.docx
  test_spelling_typeC.docx
  test_spelling_clean.docx
"""
import os
import sys

# Make sure we can import from parent dir
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import lxml.etree as etree

TEMP_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "_temp")
os.makedirs(TEMP_DIR, exist_ok=True)


def set_rtl(paragraph):
    """Force RTL on a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.insert(0, bidi)


def add_rtl_para(doc, text, style=None):
    """Add an RTL paragraph."""
    if style:
        try:
            p = doc.add_paragraph(text, style=style)
        except Exception:
            p = doc.add_paragraph(text)
    else:
        p = doc.add_paragraph(text)
    set_rtl(p)
    return p


def add_table_row(table, cells):
    """Add a row to a table."""
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = val
    return row


def create_bordered_cover_table(doc, rows_data):
    """Create a bordered cover table with 2 columns."""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    return table


# ──────────────────────────────────────────────────────────────────────────────
# DOCUMENT A — שומת נכס מקרקעין (Standard Appraisal)
# ──────────────────────────────────────────────────────────────────────────────
def create_type_a(output_path: str):
    doc = Document()

    # Cover page with bordered table
    doc.add_heading("חוות דעת שמאית", 0)
    cover_data = [
        ("גוש", "6623"),
        ("חלקה", "458"),
        ("תת חלקה", "2"),
        ("כתובת", "רחוב אשכנזי 80, תל אביב"),
        ("מזמין השומה", "הוועדה המקומית תל אביב"),
        ("מוכן על ידי", "פנחס פרסר"),          # T2: proper name — trap
        ("מספר תיק", "12005-2025"),
        ("תאריך", "01/03/2026"),
    ]
    create_bordered_cover_table(doc, cover_data)
    doc.add_page_break()

    # Para 1 — Edge case E1: empty paragraph
    doc.add_paragraph("   ")   # whitespace-only

    # Para 2 — normal
    add_rtl_para(doc, "חוות דעת שמאית זו הוכנה לפי בקשת הוועדה המקומית תל אביב.")

    # Para 3 — Edge case E2: number only
    doc.add_paragraph("120")

    # Section 1
    add_rtl_para(doc, "1. מטרת חוות הדעת", style="Heading 2")

    # Para ~4
    add_rtl_para(doc, "מטרת חוות הדעת הינה לקבוע את שווי השוק של הנכס הנדון לצורך הגשה לוועדה המקומית.")

    # Para ~5 — A1: שמאיות (should be שמאות)
    add_rtl_para(doc, "חוות הדעת הוכנה על ידי שמאיות מקרקעין מוסמך, בהתאם לתקן מספר 19.")

    # Para ~6 — E3: Hebrew+English mix — edge case
    add_rtl_para(doc, "הנכס מסווג לפי תקן IFRS 16 של ה-IASB ואינו כפוף להפחתה שנתית.")

    # Section 2
    add_rtl_para(doc, "2. תיאור הנכס", style="Heading 2")

    # Para ~7 — T4: plan reference (trap) — do NOT flag
    add_rtl_para(doc, "הנכס מיועד למגורים בהתאם לתכנית מס' תא/2834 החלה על המקרקעין.")

    # Para ~8 — A2: gender mismatch ממוקמת
    add_rtl_para(doc, "הנכס ממוקמת ברחוב אשכנזי 80, קומה 4, בלב שכונת נווה צדק בתל אביב.")

    # Para ~9 — A9: missing period at end
    add_rtl_para(doc, "הנכס כולל 4 חדרים, סלון, מטבח, חדר שירותים ומרפסת שמש")

    # Para ~10 — T5: perfectly clean paragraph (trap)
    add_rtl_para(doc, "הנכס רשום בלשכת רישום המקרקעין, גוש 6623, חלקה 458, תת חלקה 2, כדירת מגורים.")

    # Section 3
    add_rtl_para(doc, "3. מצב משפטי", style="Heading 2")

    # Para ~11 — A3: number mismatch מראה
    add_rtl_para(doc, "הנתונים מראה כי הנכס רשום על שם מר ישראל ישראלי, ללא שעבודים או עיקולים.")

    # Para ~12 — A14: wrong preposition עליו (referring to feminine — הקרקע)
    add_rtl_para(doc, "הקרקע שייכת למדינת ישראל, וזכויות הבעלים מוגדרות עליו בנסח הטאבו.")

    # Para ~13 — T9: empty field _____ (trap)
    add_rtl_para(doc, "שם בעל הזכויות הקודם: _____")

    # Para ~14 — A4: extra letter היווון (should be היוון)
    add_rtl_para(doc, "שיעור ההיווון שיושם בחישוב עמד על 7.5% בהתאם לנתוני השוק.")

    # Para ~15 — A10: missing comma לתאריך הקובע ולא
    add_rtl_para(doc, "שווי הנכס נקבע לתאריך הקובע ולא לתאריך הדוח הסופי.")

    # Section 4
    add_rtl_para(doc, "4. מצב תכנוני", style="Heading 2")

    # Para ~16 — T6: run-on sentence (phrasing issue, not spelling)
    add_rtl_para(doc,
        "הנכס ממוקם באזור מגורים מסוג א' שבו מותרות בניה לגובה של עד 4 קומות ו-25 יחידות דיור "
        "לדונם וייעוד הקרקע בתכנית המתאר הקיימת תואם את השימוש בפועל כפי שנצפה בסיור שנערך בנכס "
        "ביום 15.02.2026 ואושר על ידי השמאי ובא כוחו."
    )

    # Para ~17 — A5: doubled letter משותתף (should be משותף)
    add_rtl_para(doc, "הנכס מהווה חלק מרכוש משותתף הכולל את חדר המדרגות, החצר ומחסן.")

    # Para ~18 — normal
    add_rtl_para(doc, "שטח הרכוש המשותף עומד על 18% מסך שטח הבניין.")

    # Para ~19 — A11: unclosed paren
    add_rtl_para(doc, "פרטים נוספים ניתן למצוא בנסח הטאבו (ראה נספח א.")

    # Para ~20 — A6: missing yod בנייה
    add_rtl_para(doc, "היתר הבניה ניתן ביום 12.06.2019 למספר קומות ולשטח כפי שמוגדר בתכנית.")

    # Section 5
    add_rtl_para(doc, "5. תיאור הסביבה", style="Heading 2")

    # Para ~21 — T7: area contradiction (120 vs 85) — logic issue (trap)
    add_rtl_para(doc, "שטח הדירה עומד על 120 מ\"ר לפי הצהרת הבעלים ובדיקות שנערכו בשטח.")

    # Para ~22 — A12: space before comma
    add_rtl_para(doc, "הנכס , הממוקם בסמוך לפארק הצפוני, נהנה מנגישות מצוינת לתחבורה ציבורית.")

    # Para ~23 — A7: missing letter זכיות (should be זכויות)
    add_rtl_para(doc, "זכיות הבעלות רשומות בלשכת רישום המקרקעין בתל אביב ואינן שנויות במחלוקת.")

    # Para ~24 — T8: legal boilerplate (trap)
    add_rtl_para(doc,
        "הגבלת אחריות: חוות דעת זו הוכנה אך ורק לצורך שצוין לעיל. כל שימוש אחר או העברה לצד שלישי "
        "אסורים ללא אישור בכתב מהשמאי. לא תישא השמאית בכל אחריות לנזק שייגרם משימוש לא מורשה."
    )

    # Para ~25 — A13: missing colon שטח הנכס 120 מ"ר
    add_rtl_para(doc, "שטח הנכס 120 מ\"ר הכולל את שטח המרפסת והחנייה הצמודה.")

    # Section 6
    add_rtl_para(doc, "6. נתונים השוואתיים", style="Heading 2")

    # Para ~26 — A8: double space (two spaces between הנכס and נמצא)
    add_rtl_para(doc, "הנכס  נמצא בסמוך לשלוש עסקאות השוואה שנמכרו בשנת 2025.")

    # Comparable sales table (with error in cell — E4)
    table2 = doc.add_table(rows=4, cols=4)
    table2.style = "Table Grid"
    headers = table2.rows[0]
    headers.cells[0].text = "כתובת"
    headers.cells[1].text = "שטחי"    # E4: spelling error in table (should be שטחים)
    headers.cells[2].text = "מחיר"
    headers.cells[3].text = "תאריך"
    add_table_row(table2, ["רחוב הרצל 10", "90 מ\"ר", "2,200,000 ₪", "03/2025"])
    add_table_row(table2, ["רחוב אלנבי 45", "105 מ\"ר", "2,600,000 ₪", "07/2025"])
    add_table_row(table2, ["שדרות רוטשילד 8", "85 מ\"ר", "2,100,000 ₪", "11/2025"])

    # Section 7
    add_rtl_para(doc, "7. הערכת שווי", style="Heading 2")

    # Para ~27 — normal
    add_rtl_para(doc, "הערכת השווי בוצעה בשיטת ההשוואה הישירה בשילוב עם שיטת ההיוון.")

    # Para ~28 — A15: sofit error הנכסימ (should be הנכסים)
    add_rtl_para(doc, "לאחר השוואה לנתוני השוק, הנכסימ הדומים נמכרו בטווח של 22,000–25,000 ₪ למ\"ר.")

    # Para ~29 — E5: short correct paragraph
    add_rtl_para(doc, "ראה לעיל.")

    # Section 8
    add_rtl_para(doc, "8. סיכום", style="Heading 2")

    # Para ~30
    add_rtl_para(doc, "לאור כל האמור לעיל, שווי השוק של הנכס נקבע לסך של 2,400,000 ₪ (שניים וארבעה מאות אלף שקלים).")

    # Add header with spelling error E6: שמאות מקרקעיין
    # We inject it via python-docx header
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "סויצקי רמי — שמאות מקרקעיין"
    set_rtl(hp)

    # Footer with page numbers (clean — E7)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "סויצקי רמי שמאות מקרקעין — עמוד 1"
    set_rtl(fp)

    doc.save(output_path)
    print(f"Created: {output_path}")


# ──────────────────────────────────────────────────────────────────────────────
# DOCUMENT B — היטל השבחה (Betterment Levy)
# ──────────────────────────────────────────────────────────────────────────────
def create_type_b(output_path: str):
    doc = Document()

    doc.add_heading("שומת היטל השבחה", 0)
    cover_data = [
        ("גוש", "6854"),
        ("חלקה", "41"),
        ("תת חלקה", "2"),
        ("כתובת", "סלע 1 כניסה ב'"),
        ("שכונה", "ורדיה"),
        ("עיר", "שוהם"),
        ("מזמין השומה", "הוועדה המקומית שוהם"),
        ("המבקשים", "יוחאי פנחס פרסר"),    # T10 proper name trap
        ("מספר תיק", "2026-00330"),
        ("תאריך", "04/03/2026"),
    ]
    create_bordered_cover_table(doc, cover_data)
    doc.add_page_break()

    # Para 1
    add_rtl_para(doc, "שומה זו עוסקת בחישוב היטל השבחה בגין תכנית משביחה החלה על המקרקעין הנ\"ל.")

    # Para 2
    add_rtl_para(doc, "ההשבחה נובעת ממתן הקלה בקו הבניין ותוספת זכויות בניה על פי תכנית מאושרת.")

    # Para 3
    add_rtl_para(doc, "הנכס נרכש בשנת 2018 ומאז לא חלו שינויים במצבו הפיזי.")

    # Para ~4 — B1: smikut error בית של הספר
    add_rtl_para(doc, "ליד הנכס שוכן בית של הספר האזורי ופארק ציבורי, מה שמשפר את ערכו.")

    # Para ~5 — normal (T10 context — neighborhood name)
    add_rtl_para(doc, "שכונת \"ורדיה\" מאופיינת בבנייה נמוכה של 2-3 קומות עם גנים פרטיים.")

    # Para ~6 — normal
    add_rtl_para(doc, "מצב הדרכים הסמוכות תקין ושלם, עם תשתיות מים, חשמל וביוב.")

    # Para ~7 — B2: wrong כ/ח — כישוב instead of חישוב
    add_rtl_para(doc, "כישוב ההשבחה בוצע לפי שיטת לפני-אחרי הנהוגה בדין הישראלי.")

    # Para ~8 — B8: run-together words בהתאםלתכנית
    add_rtl_para(doc, "הזכויות ניתנו בהתאםלתכנית המשביחה שאושרה ביום 15.01.2024.")

    # Para ~9 — T13: plan reference (trap)
    add_rtl_para(doc, "תכנית מס' שד/1234 הינה התכנית המשביחה הרלוונטית לחישוב ההשבחה.")

    # Para ~10 — B3: plural form error התכניות המשביח
    add_rtl_para(doc, "התכניות המשביח נסקרו בחלק ג' של דוח זה ומפורטות אחת לאחת.")

    # Para ~11 — normal
    add_rtl_para(doc, "שווי המגרש לפני ההשבחה נקבע לפי ניתוח עסקאות השוואה בסביבה הקרובה.")

    # Para ~12 — T12: calculation numbers (trap)
    add_rtl_para(doc, "סה\"כ ההשבחה: 1,500,000 - 800,000 = 700,000 ₪ לפני מחצית.")

    # Para ~13 — B4: ש/ס swap שביבה instead of סביבה
    add_rtl_para(doc, "שביבה הנכס אופיינה בביקוש גבוה ועלייה עקבית במחירים בשנים האחרונות.")

    # Para ~14 — normal
    add_rtl_para(doc, "הנכס מהווה יחידת מגורים עצמאית הכוללת כניסה נפרדת ומרפסת פרטית.")

    # Para ~15 — B9: homophone מצא (found) instead of מצב (state)
    add_rtl_para(doc, "מצא הנכס הפיזי הינו טוב ומשמש למגורים ללא כל מגבלה תכנונית.")

    # Para ~16 — B5: missing maqaf תל אביב יפו
    add_rtl_para(doc, "עסקאות ההשוואה נלקחו מאזורים כגון שוהם, רמלה, ותל אביב יפו.")

    # Para ~17 — normal
    add_rtl_para(doc, "שיעור ההיוון לנכסים מסוג זה עומד על 5.5%-6% לפי הנתונים המעודכנים.")

    # Para ~18 — normal
    add_rtl_para(doc, "הסיור בנכס נערך ביום 28.02.2026 בנוכחות הבעלים ובא כוחם.")

    # Para ~19 — B6: wrong gershayim תב'ע (should be תב"ע)
    add_rtl_para(doc, "בהתאם לתב'ע החלה על המקרקעין, מותרת הקמת 4 יחידות דיור נוספות.")

    # Para ~20 — normal
    add_rtl_para(doc, "הממצאים הסטטיסטיים מעידים על עלייה של 12% בממוצע מחירי הדירות.")

    # Para ~21 — normal
    add_rtl_para(doc, "לאחר ניכוי עלויות הבניה, הרווח הנדרש ומיסים, שווי הקרקע לאחר ההשבחה הוא 1,500,000 ₪.")

    # Para ~22 — B7: missing space after period הנכס.הממוקם
    add_rtl_para(doc, "הנכס.הממוקם בשכונת ורדיה נהנה מסביבה ירוקה ושקטה.")

    # Para ~23 — normal
    add_rtl_para(doc, "סה\"כ היטל ההשבחה עומד על 350,000 ₪ (שלוש מאות וחמישים אלף שקלים).")

    # Para ~24 — T11: address data trap
    add_rtl_para(doc, "כתובת הנכס: סלע 1 כניסה ב', שוהם — אינה כוללת מיקוד.")

    # Para ~25 — B10: semicolon misuse
    add_rtl_para(doc, "הנכס; ממוקם בגבול המערבי של השכונה ומהווה חלק מפרויקט פיתוח.")

    # Para ~26 — summary
    add_rtl_para(doc, "בכפוף לאמור לעיל, שווי ההשבחה לצורך חישוב היטל ההשבחה עומד על 700,000 ₪.")

    doc.save(output_path)
    print(f"Created: {output_path}")


# ──────────────────────────────────────────────────────────────────────────────
# DOCUMENT C — תיקון שומה (Correction) — Hardest
# ──────────────────────────────────────────────────────────────────────────────
def create_type_c(output_path: str):
    doc = Document()

    doc.add_heading("תיקון שומה", 0)
    cover_data = [
        ("מספר תיק מקורי", "12005-2025"),
        ("מסמך", "תיקון שומה"),
        ("שמאי", "סויצקי רמי"),
        ("תאריך תיקון", "04/03/2026"),
    ]
    create_bordered_cover_table(doc, cover_data)
    doc.add_page_break()

    # Para 1
    add_rtl_para(doc, "מסמך זה מהווה תיקון לשומה המקורית שהוגשה ביום 01.01.2026.")

    # Para 2
    add_rtl_para(doc, "התיקון נערך בעקבות בקשה מטעם הוועדה המקומית לעדכן את נתוני הנכס.")

    # Para ~3 — C1: extra space inside word זכו יות
    add_rtl_para(doc, "זכו יות הבעלות בנכס הנדון עברו לבעלים הנוכחי בשנת 2020.")

    # Para ~4 — normal
    add_rtl_para(doc, "הנכס ממוקם ברחוב הרצל 10, תל אביב, קומה 3, שטח 95 מ\"ר.")

    # Para ~5 — normal
    add_rtl_para(doc, "מצב הנכס הפיזי לא השתנה מאז השומה המקורית ואינו מחייב עדכון.")

    # Para ~6 — C2: אם/עם confusion
    add_rtl_para(doc, "אם התיקון אושר, יש לעדכן את כל ההפניות לשומה המקורית בהתאם.")

    # Para ~7 — normal
    add_rtl_para(doc, "הנתונים ההשוואתיים עודכנו לפי עסקאות מהרבעון הרביעי של 2025.")

    # Para ~8 — normal
    add_rtl_para(doc, "שיעור ההיוון לא השתנה ועומד על 7.5% כפי שנקבע בשומה המקורית.")

    # Para ~9 — C3: missing he hayedia — נכס ממוקם ב (should be הנכס)
    add_rtl_para(doc, "נכס ממוקם ב-תל אביב ומסווג כנכס מגורים לפי התכנית החלה עליו.")

    # Para ~10 — C7: אחר/אחד ambiguous
    add_rtl_para(doc, "שני נכסים מגרש בסמיכות: הנכס האחר נרכש ב-2022, המשמש כמחסן.")

    # Para ~11 — normal
    add_rtl_para(doc, "הערכת השווי בוצעה בשיטת ההשוואה הישירה ללא שינוי מהשומה המקורית.")

    # Para ~12 — C4: CRITICAL TRAP — שומה הוגשה is CORRECT (feminine)
    add_rtl_para(doc, "השומה הוגשה בהתאם לדרישות החוק ואושרה על ידי מחלקת השמאות.")

    # Para ~13 — normal
    add_rtl_para(doc, "אין שינויים מהותיים בנתוני שוק שיש בהם כדי לשנות את קביעת השווי.")

    # Para ~14 — normal
    add_rtl_para(doc, "הנכס ממשיך להיות מוגדר כנכס מגורים ב-4 חדרים, ללא שינוי בייעוד.")

    # Para ~15 — C5: consecutive dots הנכס.. ממוקם
    add_rtl_para(doc, "הנכס.. ממוקם בסמוך לתחנת רכבת תל אביב ומרכז עסקים.")

    # Para ~16 — normal
    add_rtl_para(doc, "לאחר עיון בכל הנתונים, ניתן לאשר את קביעת השווי המקורית.")

    # Para ~17 — normal
    add_rtl_para(doc, "שווי הנכס לתאריך הקובע: 2,400,000 ₪, ללא שינוי משומת הבסיס.")

    # Para ~18 — C6: missing space before paren הנכס(ראה נספח)
    add_rtl_para(doc, "הנכס(ראה נספח ב) מהווה יחידת מגורים עצמאית הרשומה כתת חלקה.")

    # Para ~19 — closing
    add_rtl_para(doc, "תיקון שומה זה מחליף את הגרסה הקודמת לכל עניין ודבר.")

    # Para ~20 — normal
    add_rtl_para(doc, "השמאי מאשר כי הנתונים שבדוח זה נכונים ומדויקים למיטב ידיעתו.")

    doc.save(output_path)
    print(f"Created: {output_path}")


# ──────────────────────────────────────────────────────────────────────────────
# DOCUMENT CLEAN — 15 paragraphs of perfect Hebrew
# ──────────────────────────────────────────────────────────────────────────────
def create_clean(output_path: str):
    doc = Document()

    doc.add_heading("חוות דעת שמאית — מסמך בדיקה", 0)

    clean_paras = [
        "חוות דעת שמאית זו נערכה לבקשת הוועדה המקומית לשם קביעת שווי שוק.",
        "הנכס הנדון ממוקם ברחוב הרצל 10, תל אביב, גוש 6623, חלקה 458, תת חלקה 2.",
        "הנכס רשום בלשכת רישום המקרקעין ומסווג כדירת מגורים בת 4 חדרים.",
        "שטח הנכס עומד על 95 מ\"ר, הכולל מרפסת שמש בשטח 12 מ\"ר.",
        "הנכס נרכש בשנת 2015 ונמצא במצב תחזוקה טוב.",
        "הסיור בנכס נערך ביום 28.02.2026 בנוכחות הבעלים.",
        "הנכס מיועד למגורים בהתאם לתכנית מתאר מס' תא/3000 החלה על המקרקעין.",
        "נסח הטאבו מיום 01.03.2026 מאשר כי אין שעבודים או עיקולים הרשומים על הנכס.",
        "הערכת השווי בוצעה בשיטת ההשוואה הישירה, בהתייחס לשלוש עסקאות השוואה.",
        "עסקאות ההשוואה נמכרו בטווח של 23,000–26,000 ₪ למ\"ר בשנת 2025.",
        "לאחר ניתוח הנתונים ובהתחשב בפרמטרים הרלוונטיים, נקבע שווי השוק.",
        "שווי השוק של הנכס לתאריך הקובע הינו 2,500,000 ₪ (שניים וחצי מיליון שקלים).",
        "חוות דעת זו נערכה בהתאם לתקן שמאות מספר 19 של מועצת שמאי המקרקעין.",
        "השמאי מצהיר כי אין לו עניין אישי בנכס ואין לו ניגוד עניינים.",
        "חוות דעת זו מיועדת אך ורק למטרה שצוינה לעיל ואינה מיועדת לכל שימוש אחר.",
    ]

    for para in clean_paras:
        add_rtl_para(doc, para)

    doc.save(output_path)
    print(f"Created: {output_path}")


# ──────────────────────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    paths = {
        "A": os.path.join(TEMP_DIR, "test_spelling_typeA.docx"),
        "B": os.path.join(TEMP_DIR, "test_spelling_typeB.docx"),
        "C": os.path.join(TEMP_DIR, "test_spelling_typeC.docx"),
        "clean": os.path.join(TEMP_DIR, "test_spelling_clean.docx"),
    }

    print("=== Generating test documents ===")
    create_type_a(paths["A"])
    create_type_b(paths["B"])
    create_type_c(paths["C"])
    create_clean(paths["clean"])

    print("\n=== All documents created ===")
    for k, p in paths.items():
        size = os.path.getsize(p)
        print(f"  {k}: {p} ({size:,} bytes)")
