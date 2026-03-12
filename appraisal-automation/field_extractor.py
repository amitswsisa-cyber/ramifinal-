"""
field_extractor.py
Dynamic extraction of label:value pairs from Hebrew DOCX cover pages.

Revised strategy (v2):
 1. Read paragraph text from the first ~10 paragraphs (all real docs use paragraphs
    for the cover box, not tables).
 2. For each non-empty line:
    a. Split on ALL colon-separated label:value pairs — handles multiple pairs per line
       e.g. 'גוש : 6623   חלקה: 458' → {גוש:6623, חלקה:458}
       e.g. 'גוש: 6854 חלקה: 41 תת חלקה: 2' → {גוש:6854, חלקה:41, תת חלקה:2}
    b. For lines with no colon, classify by content:
       - Starts with רחוב / street pattern → 'רחוב'
       - Starts with הסלע / סלע           → 'סלע/כניסה'
       - Contains שכונת / שכונה           → 'שכונה'  (value = quoted text or full line)
       - Short word(s), no digits, not yet matched → 'עיר'
 3. Also fallback-scan tables (for future doc types that use tables on cover).

Does NOT use python-docx for XML editing — only for reading text.
"""
import io
import re
from docx import Document

# Hebrew Unicode block — U+0590–U+05FF
_HE = r'\u0590-\u05FF'

# A "Hebrew word" is 1+ Hebrew chars (we include spaces so multi-word labels work)
# Label ends with optional spaces + colon (ASCII ':' — '׃' is the Hebrew punctuation
# mark but these docs use ASCII).
#
# Regex: find every position in a line that looks like  <label> :
# Where label = sequence of Hebrew chars/spaces NOT containing another colon.
_LABEL_RE = re.compile(
    r'([' + _HE + r'][' + _HE + r'\s]*?)\s*:'  # label (Hebrew words) followed by colon
)


# ── Public API ────────────────────────────────────────────────────────────────

def extract_cover_fields(file_obj) -> dict[str, str]:
    """
    Accept a file-like object (Streamlit UploadedFile or open() file).
    Return {label: value} for all pairs found on the cover page.
    """
    doc = _load_doc(file_obj)
    fields: dict[str, str] = {}

    # ── Pass 1: paragraph scan (primary source for these document types) ──────
    cover_paras = _get_cover_paragraphs(doc)

    unmatched_lines: list[str] = []

    for line in cover_paras:
        line = line.strip()
        if not line:
            continue

        pairs = _extract_pairs_from_line(line)
        if pairs:
            for label, value in pairs.items():
                if label and label not in fields:
                    fields[label] = value
        else:
            # No colon found — might be a street/city/neighborhood/sela line
            unmatched_lines.append(line)

    # ── Pass 2: classify label-less lines ────────────────────────────────────
    for line in unmatched_lines:
        label, value = _classify_labelless_line(line, fields)
        if label and label not in fields:
            fields[label] = value

    # ── Pass 3: table scan (fallback for documents that use cover table) ─────
    for table in doc.tables[:5]:
        _scan_table(table, fields)

    # ── Pass 4: extract body fields (section 6 table-style + Table 1) ─────
    _extract_body_fields(doc, fields)

    return fields


def detect_document_type(file_obj) -> str:
    """
    Detect document type from the title on page 1.
    Returns: "betterment" | "correction" | "standard"
    """
    doc = _load_doc(file_obj)

    texts = []
    for para in doc.paragraphs[:20]:
        texts.append(para.text)
    for table in doc.tables[:2]:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)

    combined = " ".join(texts)

    if "היטל השבחה" in combined:
        return "betterment"
    if "תיקון שומה" in combined:
        return "correction"
    return "standard"


# ── Internal helpers ──────────────────────────────────────────────────────────

def _load_doc(file_obj) -> Document:
    """Load a DOCX from a file-like object or path, resetting seek position."""
    if hasattr(file_obj, "read"):
        data = file_obj.read()
        if hasattr(file_obj, "seek"):
            file_obj.seek(0)
        return Document(io.BytesIO(data))
    return Document(file_obj)


def _get_cover_paragraphs(doc: Document) -> list[str]:
    """
    Return text lines from the first meaningful paragraphs.
    Stop at 'תוכן ענינים' or after 15 non-empty paragraphs — whichever comes first.
    """
    lines: list[str] = []
    count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        # Stop at table-of-contents marker
        if "תוכן ענינים" in text or "ת ו כ ן" in text:
            break
        if text:
            lines.append(text)
            count += 1
        if count >= 15:
            break
    return lines


def _extract_pairs_from_line(line: str) -> dict[str, str]:
    """
    Extract ALL label:value pairs from a single line.

    Handles:
      'גוש : 6623   חלקה: 458'          → {גוש: 6623, חלקה: 458}
      'גוש: 6854 חלקה: 41 תת חלקה: 2'  → {גוש: 6854, חלקה: 41, תת חלקה: 2}
      'מזמין השומה: הועדה המקומית שוהם' → {מזמין השומה: הועדה המקומית שוהם}
      'מספר תיק : 2025-12005'           → {מספר תיק: 2025-12005}

    Returns empty dict if no colon found in the line.
    """
    if ":" not in line:
        return {}

    pairs: dict[str, str] = {}

    # Find all label positions using the regex
    # Each match gives us: (label_text, match.end()) — start of value
    matches = list(_LABEL_RE.finditer(line))
    if not matches:
        return {}

    for i, match in enumerate(matches):
        label = match.group(1).strip()
        value_start = match.end()

        # Value runs until the START of the next label (or end of line)
        if i + 1 < len(matches):
            value_end = matches[i + 1].start()
        else:
            value_end = len(line)

        value = line[value_start:value_end].strip()

        # Clean up: remove trailing separators like spaces, slashes, dashes
        value = value.rstrip(" /\t")

        if label:
            pairs[label] = value

    return pairs


# Known document title/header keywords — lines containing these should NOT be
# classified as street or city even if they look like short Hebrew phrases.
_TITLE_KEYWORDS = (
    "היטל השבחה",
    "שומת נכס",
    "תיקון שומה",
    "חוות דעת",
    "דוח שמאות",
    "ת ו כ ן",
    "מדור היטל",
    "מנהל ההנדסה",
    "הועדה המקומית לתכנון ובניה",
    "לכבוד",
    "הגבלת שימוש",
)


def _is_title_line(line: str) -> bool:
    """Return True if the line is a document title/header, not a data field."""
    for kw in _TITLE_KEYWORDS:
        if kw in line:
            return True
    # Spaced-out title chars: 'ת ו כ ן    ע נ י נ י ם'
    if re.search(r'.{1}\s.{1}\s.{1}\s.{1}', line):
        return True
    return False


def _classify_labelless_line(line: str, existing_fields: dict[str, str]) -> tuple[str, str]:
    """
    Classify a line that has no colon into a known field.
    Returns (label, value) or ('', '') if it can't be classified.
    """
    # ── Special case: title field with dash separator (BEFORE title-skip) ────
    # Pattern: 'שומת נכס מקרקעין מלאה - דירת מגורים' or 'שומת נכס מקרקעין מלאה - ________'
    # Must be checked BEFORE _is_title_line() which would otherwise discard it.
    if "שומת נכס" in line and "-" in line:
        parts = line.rsplit("-", 1)
        if len(parts) == 2:
            value = parts[1].strip()
            if value and "סוג שומה" not in existing_fields:
                return "סוג שומה", value

    # Skip document title / header lines
    if _is_title_line(line):
        return "", ""

    # Skip lines with numbering (table-of-contents entries: '1) מטרת...', '2. ב)')
    if re.match(r'^\d+[\)\.]', line):
        return "", ""

    words = line.split()

    # 1. Street: starts explicitly with the word 'רחוב'
    if line.startswith("רחוב"):
        return "רחוב", line[len("רחוב"):].strip()

    # 2. Sela/entrance: starts with הסלע or סלע — check BEFORE digit heuristic
    #    so 'הסלע 1 כניסה ב' is not misclassified as a street address.
    if line.startswith("הסלע") or line.startswith("סלע"):
        return "סלע/כניסה", line

    # 3. Street heuristic: Hebrew words + at least one digit → likely a street address.
    #    e.g. 'אשכנזי 80', 'פתחיה מרגנשבורג 41'
    if (
        any(c.isdigit() for c in line)
        and len(words) <= 6
        and _is_mostly_hebrew(line)
        and "רחוב" not in existing_fields
        and "גוש" not in line
        and "חלקה" not in line
    ):
        return "רחוב", line

    # 4. Neighborhood: contains שכונת or שכונה
    if "שכונת" in line or "שכונה" in line:
        # Extract quoted text if present: שכונת המגורים "ורדים" → ורדים
        quoted = re.search(r'["\u201c\u201d]([^"\u201c\u201d]+)["\u201c\u201d]', line)
        if quoted:
            value = quoted.group(1).strip()
        else:
            value = line
            for prefix in ("שכונת המגורים", "שכונת", "שכונה"):
                value = value.replace(prefix, "").strip()
        return "שכונה", value

    # 5. City: short 1-3 Hebrew words, no digits, not already matched
    if (
        1 <= len(words) <= 3
        and _is_mostly_hebrew(line)
        and not any(c.isdigit() for c in line)
        and "עיר" not in existing_fields
    ):
        return "עיר", line

    return "", ""


def _is_mostly_hebrew(text: str) -> bool:
    """Return True if the majority of alphabetic characters are Hebrew."""
    he_chars = sum(1 for c in text if '\u0590' <= c <= '\u05FF')
    alpha_chars = sum(1 for c in text if c.isalpha())
    return alpha_chars > 0 and he_chars / alpha_chars > 0.6


def _scan_table(table, fields: dict[str, str]) -> None:
    """
    Fallback: scan a table for label:value pairs.
    Mutates `fields` in-place.
    """
    for row in table.rows:
        cells = row.cells
        n = len(cells)
        i = 0
        while i < n:
            text = cells[i].text.strip()
            if not text:
                i += 1
                continue

            pairs = _extract_pairs_from_line(text)
            if pairs:
                for label, value in pairs.items():
                    # Skip disclaimer / boilerplate lines (very long values)
                    if label and label not in fields and len(value) <= 120:
                        fields[label] = value
                i += 1
                continue

            # Adjacent cell pattern: "label:" | "value"
            if text.endswith(":"):
                label = text.rstrip(":").strip()
                if i + 1 < n:
                    value = cells[i + 1].text.strip()
                    if label and label not in fields:
                        fields[label] = value
                    i += 2
                    continue

            i += 1


# ── Body field extraction (section 6 + Table 1) ──────────────────────────────

# Labels to extract from body paragraphs (tab-separated or colon-separated)
_BODY_FIELD_LABELS = {"שטח חלקה", "שטח בנוי", "תיאור זכויות"}

# Labels to extract from Table 1 (פרטי הנכס)
_TABLE1_LABELS_SET = {"גוש", "חלקה", "שטח חלקה", "תת חלקה", "תת-חלקה",
                      "שטח בנוי", "תיאור זכויות", "מיקום", "החלקה הנישום", "מגרש"}


def _extract_body_fields(doc, fields: dict[str, str]) -> None:
    """
    Extract additional fields from the document body that aren't on the cover page.
    These come from:
    1. Section 6 tab-separated lines
    2. Table 1 (פרטי הנכס) — table rows with known labels
    """
    # Scan body paragraphs for colon-separated pairs with body-specific labels
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        pairs = _extract_pairs_from_line(text)
        for label, value in pairs.items():
            if label in _BODY_FIELD_LABELS and label not in fields:
                fields[label] = value

    # Scan all tables for Table 1 labels
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            for i, cell in enumerate(cells):
                cell_text = cell.text.strip()
                label_candidate = cell_text.rstrip(":").strip()
                if label_candidate in _TABLE1_LABELS_SET:
                    field_name = label_candidate.replace("-", " ")
                    if field_name in fields:
                        continue
                    for offset in [-1, 1]:
                        adj_idx = i + offset
                        if 0 <= adj_idx < len(cells):
                            val = cells[adj_idx].text.strip()
                            if val and val != cell_text and not re.fullmatch(r'_+', val):
                                fields[field_name] = val
                                break
